from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION_START
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

OUT = Path("docs/Historical_Battle_of_Tukayyid_Campaign_Rules.docx")
# compact_reference_guide with a named Core-inspired manual override:
# tighter page geometry, compact body type, steel-blue hierarchy, and two-column Track text.
BLUE, DARK, PALE, WHITE, GREY = "607FAA", "213854", "DCE3EE", "FFFFFF", "566273"

def shade(cell, fill):
    pr = cell._tc.get_or_add_tcPr(); el = OxmlElement("w:shd"); el.set(qn("w:fill"), fill); pr.append(el)

def cell_margins(cell, top=80, start=120, bottom=80, end=120):
    pr = cell._tc.get_or_add_tcPr(); mar = pr.first_child_found_in("w:tcMar")
    if mar is None: mar = OxmlElement("w:tcMar"); pr.append(mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        x = OxmlElement(f"w:{name}"); x.set(qn("w:w"), str(value)); x.set(qn("w:type"), "dxa"); mar.append(x)

def keep(p):
    p.paragraph_format.keep_with_next = True
    return p

def set_columns(section, count=2, space=360):
    sectPr = section._sectPr
    cols = sectPr.xpath("./w:cols")
    el = cols[0] if cols else OxmlElement("w:cols")
    if not cols: sectPr.append(el)
    el.set(qn("w:num"), str(count)); el.set(qn("w:space"), str(space)); el.set(qn("w:equalWidth"), "1")

def chapter_band(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10); p.paragraph_format.space_after = Pt(8); p.paragraph_format.keep_with_next = True
    pr = p._p.get_or_add_pPr(); shd = OxmlElement("w:shd"); shd.set(qn("w:fill"), DARK); pr.append(shd)
    r = p.add_run(text.upper()); r.bold = True; r.font.name = "Aptos Display"; r.font.size = Pt(15); r.font.color.rgb = RGBColor.from_string(WHITE)
    return p

def bullet(text, level=0):
    p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    p.add_run(text); return p

def rule(label, text):
    p = doc.add_paragraph(); p.add_run(label + ". ").bold = True; p.add_run(text); return p

def historical_excerpt():
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(.18)
    p.paragraph_format.right_indent = Inches(.18)
    p.paragraph_format.space_before = Pt(7)
    p.paragraph_format.space_after = Pt(11)
    p.paragraph_format.line_spacing = 1.03
    pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd"); shd.set(qn("w:fill"), "EEF2F7"); pr.append(shd)
    borders = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    for key, value in (("val", "single"), ("sz", "20"), ("space", "8"), ("color", BLUE)):
        left.set(qn(f"w:{key}"), value)
    borders.append(left); pr.append(borders)

    passages = [
        "In May 3052, the armies of ComStar met the invading Clans on Tukayyid. For twenty-one days, they fought across seven campaign regions to decide whether the Clan advance toward Terra would continue. ComStar prevailed, and the resulting truce halted the invasion for fifteen years.",
        "At the time, Tukayyid was celebrated as the battle that saved the Inner Sphere. Its defenders had denied the Clans their ultimate prize and given the Great Houses time to recover, rearm, and prepare for the conflicts that followed.",
        "Yet, as the battle's centennial approaches, Terra rests in Clan hands and an ilClan rules from humanity's birthplace. Tukayyid unquestionably changed the course of history, but its victory came at a terrible price.",
        "Was the time purchased on Tukayyid worth its cost—and what did the Inner Sphere ultimately do with it?",
    ]
    for i, passage in enumerate(passages):
        r = p.add_run(passage); r.font.size = Pt(8.8)
        if i == len(passages) - 1: r.italic = True
        r.add_break()
        if i < len(passages) - 1: r.add_break()
    r = p.add_run("—Dr. Elara Venn, Tukayyid: A Century Later, Tukayyid Memorial Institute, 3151")
    r.bold = True; r.italic = True; r.font.size = Pt(8.5); r.font.color.rgb = RGBColor.from_string(DARK)
    return p

def table(headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers)); t.style = "Table Grid"; t.alignment = WD_TABLE_ALIGNMENT.LEFT
    t.autofit = False
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]; c.text = h; shade(c, PALE); cell_margins(c)
        c.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for r in c.paragraphs[0].runs: r.bold = True; r.font.color.rgb = RGBColor.from_string(DARK)
    trpr = t.rows[0]._tr.get_or_add_trPr(); repeat = OxmlElement("w:tblHeader"); repeat.set(qn("w:val"), "true"); trpr.append(repeat)
    for row in rows:
        cells = t.add_row().cells
        trpr = t.rows[-1]._tr.get_or_add_trPr(); cant = OxmlElement("w:cantSplit"); trpr.append(cant)
        for i, value in enumerate(row):
            cells[i].text = str(value); cell_margins(cells[i]); cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            for p in cells[i].paragraphs: p.paragraph_format.space_after = Pt(0)
    total_dxa = 10040
    if widths:
        scale = total_dxa / sum(widths)
        col_dxa = [round(w * scale) for w in widths]
        col_dxa[-1] += total_dxa - sum(col_dxa)
    else:
        col_dxa = [total_dxa // len(headers)] * len(headers)
        col_dxa[-1] += total_dxa - sum(col_dxa)
    tblPr = t._tbl.tblPr
    tblW = tblPr.first_child_found_in("w:tblW"); tblW.set(qn("w:type"), "dxa"); tblW.set(qn("w:w"), str(total_dxa))
    tblInd = OxmlElement("w:tblInd"); tblInd.set(qn("w:type"), "dxa"); tblInd.set(qn("w:w"), "120"); tblPr.append(tblInd)
    layout = OxmlElement("w:tblLayout"); layout.set(qn("w:type"), "fixed"); tblPr.append(layout)
    grid = t._tbl.tblGrid
    for child in list(grid): grid.remove(child)
    for width in col_dxa:
        gc = OxmlElement("w:gridCol"); gc.set(qn("w:w"), str(width)); grid.append(gc)
    for row in t.rows:
        for i, width in enumerate(col_dxa):
            tcPr = row.cells[i]._tc.get_or_add_tcPr(); tcW = tcPr.first_child_found_in("w:tcW")
            tcW.set(qn("w:type"), "dxa"); tcW.set(qn("w:w"), str(width))
            row.cells[i].width = Inches(width / 1440)
    spacer = doc.add_paragraph(); spacer.paragraph_format.space_after = Pt(2)
    return t

def page_number(paragraph):
    paragraph.add_run("HISTORICAL: BATTLE OF TUKAYYID   •   ")
    fld = OxmlElement("w:fldSimple"); fld.set(qn("w:instr"), "PAGE"); paragraph._p.append(fld)

def track(name, items):
    doc.add_heading(name, level=2)
    for label, text in items: rule(label, text)

doc = Document()
sec = doc.sections[0]
sec.page_width, sec.page_height = Inches(8.5), Inches(11)
sec.top_margin = Inches(.78); sec.bottom_margin = Inches(.65)
sec.left_margin = sec.right_margin = Inches(.68)
sec.header_distance = Inches(.60); sec.footer_distance = Inches(.492)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Aptos"; normal.font.size = Pt(9.4); normal.font.color.rgb = RGBColor(28, 31, 36)
normal.paragraph_format.space_after = Pt(4); normal.paragraph_format.line_spacing = 1.08
for name, size, color, before, after in (("Heading 1",17,DARK,15,7),("Heading 2",12.5,BLUE,10,4),("Heading 3",10.5,DARK,7,3)):
    s = styles[name]; s.font.name = "Aptos Display"; s.font.size = Pt(size); s.font.bold = True; s.font.color.rgb = RGBColor.from_string(color)
    s.paragraph_format.space_before = Pt(before); s.paragraph_format.space_after = Pt(after); s.paragraph_format.keep_with_next = True
styles["Title"].font.name = "Aptos Display"; styles["Title"].font.size = Pt(28); styles["Title"].font.bold = True; styles["Title"].font.color.rgb = RGBColor.from_string(DARK)
for name in ("List Bullet", "List Bullet 2"):
    styles[name].font.name = "Aptos"; styles[name].font.size = Pt(9.4); styles[name].paragraph_format.space_after = Pt(3); styles[name].paragraph_format.line_spacing = 1.08

hdr = sec.header.paragraphs[0]; hdr.text = "HISTORICAL: BATTLE OF TUKAYYID  •  CAMPAIGN RULES"; hdr.alignment = WD_ALIGN_PARAGRAPH.RIGHT
for r in hdr.runs: r.font.size = Pt(8); r.font.bold = True; r.font.color.rgb = RGBColor.from_string(BLUE)
ftr = sec.footer.paragraphs[0]; ftr.alignment = WD_ALIGN_PARAGRAPH.CENTER; page_number(ftr)
for r in ftr.runs: r.font.size = Pt(8); r.font.color.rgb = RGBColor.from_string(GREY)

p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before = Pt(65)
r = p.add_run("HISTORICAL:"); r.bold = True; r.font.name = "Aptos Display"; r.font.size = Pt(24); r.font.color.rgb = RGBColor.from_string(BLUE)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("BATTLE OF TUKAYYID"); r.bold = True; r.font.name = "Aptos Display"; r.font.size = Pt(31); r.font.color.rgb = RGBColor.from_string(DARK)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before = Pt(12)
r = p.add_run("Version 0.1 Playtest"); r.italic = True; r.font.size = Pt(12); r.font.color.rgb = RGBColor.from_string(GREY)
doc.add_paragraph().paragraph_format.space_after = Pt(18)
p = doc.add_paragraph("This is a work-in-progress update to the Battle of Tukayyid campaign system. It retains the original campaign's structure and character while incorporating selected rules and concepts from the BattleTech Core Rulebook and Hot Spots: Draconis Reach.")
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p = doc.add_paragraph("Only changed or additional rules are presented below. Unless specifically stated otherwise, continue using the published Battle of Tukayyid and Hot Spots: Draconis Reach rules, together with either the BattleTech Core Rulebook or Total Warfare and BattleTech: Mercenaries.")
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p = doc.add_paragraph("All feedback is welcome, whether based on actual play or simply reviewing the rules. Feedback on balance, campaign pacing, clarity, Track objectives, and the campaign economy would be especially helpful.")
p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after = Pt(18)
historical_excerpt()
table(["Rules Path", "Required Publications"], [
    ["New core rules", "BattleTech Core Rulebook"],
    ["Alternative older rules", "Total Warfare and BattleTech: Mercenaries"],
    ["Required with either path", "Battle of Tukayyid and Hot Spots: Draconis Reach"],
], [2.15, 4.15])
doc.add_paragraph()
rule("Using These Rules", "This document changes, replaces, or adds to the published campaign rules. Use either the BattleTech Core Rulebook or the combination of Total Warfare and BattleTech: Mercenaries for the core game and Battlefield Support rules. Battle of Tukayyid and Hot Spots: Draconis Reach are required with either rules path. Any applicable published rule not addressed here remains in effect.")
rule("Rules Priority", "If two rules conflict, use this document. A Track rule always overrides a general campaign rule.")

doc.add_page_break()
chapter_band("1  Campaign Framework")
rule("Campaign Structure", "Retain the Battle of Tukayyid Regions, force construction, campaign sequence, and original-owner recovery framework.")
rule("Orders", "Orders objectives are not used in this campaign.")
rule("Objective Points", "Objectives award Objective Points (OP). OP determine the winner of the current Track and are then discarded; they never carry into another Track.")
rule("Strategic Points", "Strategic Points (SP) are the campaign's only persistent resource. Warchest Points are not used. The side that wins the most Regions wins the campaign.")

doc.add_heading("Campaign Scale and Battle Value", level=2)
table(["Campaign Size", "Campaign BV", "33%", "25%", "50%"], [
    ["Star", "40,000", "13,200", "10,000", "20,000"],
    ["Binary", "80,000", "26,400", "20,000", "40,000"],
    ["Trinary", "120,000", "39,600", "30,000", "60,000"],
], [1.4,1.2,1.2,1.2,1.2])
rule("Fixed Denominator", "All Track percentages use the original Campaign BV allowance shown above. Losses, repairs, and replacements do not change that denominator.")
rule("Skill Adjustment", "Use the deployed pilot's actual skills when calculating a unit's BV.")
rule("Clan Advantage", "For Campaign Force construction and Track deployment limits only, treat each Clan pilot as one Experience Rating worse when accounting for BV. Actual skills do not change. This adjustment does not reduce SP purchase or repair costs.")

doc.add_heading("Starting SP and Support Allocation", level=2)
table(["Campaign Size", "Clan Starting SP", "ComStar Starting SP"], [
    ["Star", "13,200", "26,400"], ["Binary", "26,400", "52,800"], ["Trinary", "39,600", "79,200"]
], [2.0,2.0,2.0])
doc.add_page_break()
doc.add_heading("Support Allocation", level=2)
table(["Campaign Size", "Result", "Clan", "ComStar"], [
    ["Star", "Unsuccessful / Successful / Complete", "1,100 / 2,200 / 3,300", "2,200 / 4,400 / 6,600"],
    ["Binary", "Unsuccessful / Successful / Complete", "2,200 / 4,400 / 6,600", "4,400 / 8,800 / 13,200"],
    ["Trinary", "Unsuccessful / Successful / Complete", "3,300 / 6,600 / 9,900", "6,600 / 13,200 / 19,800"],
], [1.1,2.2,2.0,2.0])
bullet("Award Support Allocation after every completed Track. Unspent SP carries forward.")
bullet("Unsuccessful: the force completed at least one objective but scored fewer Objective Points.")
bullet("Successful: the force scored more Objective Points.")
bullet("Complete Success: the force won the Track and completed every available objective.")
bullet("Tie: both forces receive the Unsuccessful allocation.")
bullet("A force that refuses a Track or withdraws without completing an objective receives no allocation.")
bullet("Only the single highest applicable tier is awarded.")

doc.add_page_break()
chapter_band("2  Between Tracks")
doc.add_heading("SP Activities", level=2)
table(["Activity", "SP Cost"], [
    ["Repair armor only", "Unit tonnage ÷ 2"], ["Repair structure or critical damage", "Unit tonnage × 2"],
    ["Repair a crippled unit", "Unit tonnage × 3"], ["Repair a destroyed, recoverable unit", "Unit tonnage × 5"],
    ["Clan or Mixed Technology repair", "Multiply repair cost by 1.5; round up"],
    ["Combat vehicle or battle armor repair", "Halve repair cost; round up"],
    ["Purchase replacement unit", "Actual pilot-adjusted BV in SP"], ["Sell fully repaired unit", "BV ÷ 2"],
    ["Scrap recoverable destroyed unit", "BV ÷ 4"], ["Standard ammunition", "10 SP per ton"],
    ["Advanced or experimental ammunition", "100 SP per ton"],
    ["Reconfigure OmniMech, OmniVehicle, or modular BA", "Unit tonnage ÷ 2"],
    ["Heal MechWarrior", "30 SP per wound; maximum one wound between Tracks"],
    ["Heal battle armor trooper", "10 SP"], ["Replace killed battle armor trooper", "20 SP"],
], [4.4,2.9])
rule("Repair Category", "Use only the highest repair category that applies to a unit. After determining that category, apply all Technology Base and unit-type modifiers, then round fractions up.")
rule("Availability", "Armor-only repairs, rearming, and reconfiguration are ready for the next Track. Structure, critical, crippled, and destroyed-unit repairs are unavailable until the next Region. Replacement units are also unavailable until the next Region.")
rule("Reconfiguration", "A unit must be fully repaired before reconfiguration.")

doc.add_heading("Replacement and Recovery", level=2)
rule("Replacement Slot", "Only a destroyed unit creates a replacement slot. The replacement must be the same broad type: BattleMech, battle armor, combat vehicle, or conventional infantry. Weight class, model, and variant may change if the replacement is legal and affordable.")
rule("Selection", "Select a replacement from the relevant faction or Tukayyid Random Assignment Table. Its pilot uses the skill level purchased for that roster slot. A replacement begins with full standard ammunition.")
rule("Recovery Roll", "After a Track, the original owner rolls 2D6 for each eligible destroyed unit. On a result of 9+, the unit is recoverable. On a result of 8 or less, it is lost. Conventional infantry, crashed aerospace units, and units destroyed by the final damage from artillery or bombing cannot be recovered.")
rule("Truly Destroyed", "A BattleMech is truly destroyed when its center-torso structure is eliminated. A vehicle is truly destroyed by a fuel-tank explosion or by elimination of internal structure in a non-turret, non-rotor location. Battle armor uses a 7+ unit-survival check. A truly destroyed unit cannot be recovered or repaired.")
rule("Ownership", "Enemy units are never captured. Every recoverable unit remains the property of its original owner.")

doc.add_heading("Personnel", level=2)
bullet("A vehicle crew is killed when its vehicle is destroyed.")
bullet("Each Commander Hit or Crew Stunned result inflicts one wound on the affected crew member.")
bullet("Each battle armor trooper survives destruction on 7+; Clan battle armor applies a −2 modifier to this target number.")

chapter_band("3  Named Commanders")
table(["Campaign Size", "Starting Named Commanders", "Maximum"], [["Star","2","4"],["Binary","4","8"],["Trinary","6","12"]], [2.4,2.4,2.0])
rule("Designation", "Each Star or Level II may designate one named commander. Named status is permanent. Unnamed personnel may act as ordinary formation leaders but gain no named benefits.")
rule("Initial Commander", "An initial named commander uses the unit's purchased skill level, begins with Edge 1, and receives 150 advancement SP. The 150 SP may buy Gunnery, Piloting, Edge, or Edge Abilities, but not Special Command Abilities.")
rule("Promotion", "An eligible unnamed pilot may be formally promoted only between Tracks, after casualties are resolved and before the next Track force is selected. A promoted commander retains current skills, begins with Edge 1, receives no free advancement SP, and becomes permanently named.")
rule("Unnamed Personnel", "Unnamed pilots do not improve skills or Edge. They may record Commendations for promotion priority.")

doc.add_heading("Advancement", level=2)
rule("Campaign Spending", "After each Track, a force may divert up to 200 SP per Campaign Scale from its Support Allocation to participating named commanders, with a maximum of 100 SP assigned to any one commander.")
rule("Commendation Bonus", "A named commander who receives that Track's Commendation gains 20 additional advancement SP.")
table(["Improvement", "New Rating / Level", "Cumulative SP"], [
    ["Gunnery", "3 / 2 / 1 / 0", "300 / 700 / 1,200 / 2,200"],
    ["Piloting", "4 / 3 / 2 / 1", "100 / 200 / 700 / 1,200"],
    ["Alpha Strike Skill", "3 / 2 / 1 / 0", "400 / 900 / 1,900 / 3,400"],
    ["Edge", "2 / 3 / 4 / 5 / 6", "60 / 120 / 200 / 300 / 420"],
    ["Edge", "7 / 8 / 9 / 10", "560 / 720 / 900 / 1,100"],
    ["Edge Abilities", "1 / 2 / 3 / 4 / 5", "60 / 180 / 360 / 600 / 900"],
], [2.1,2.7,2.6])
rule("Edge", "At the start of each Track, each named commander's Edge refreshes to its current rating; unspent Edge never carries over. After rolling, spend 1 Edge to add +1 to an attack roll, reroll a motive-damage result, or reroll a critical hit scored against that commander's unit. Edge and Edge Abilities belong to the commander and cannot be transferred.")
rule("Excluded Systems", "Do not use the Draconis Reach handicap/BSP compensation system or unmatched Special Command Ability initiative rules.")

doc.add_heading("Formation and Command", level=2)
rule("Formation Training", "Formation training is organizational and free. Assign it through the parent Level III or IV, or the applicable Clan Binary, Trinary, or Cluster, based on formation type. It may change when assignments change between Tracks. It has no SP cost or delay.")
rule("Special Command Abilities", "Retain the existing Tukayyid Combatant Special Command Abilities. No commander may buy, learn, or replace an SCA during this campaign.")
rule("Acting Leader", "If a formation commander is killed, destroyed, or withdraws, appoint any surviving pilot or unit commander in that formation as Acting Leader. The Acting Leader does not inherit Edge, Edge Abilities, or training. A named Acting Leader retains only their own benefits.")

doc.add_heading("Commendations and Retirement", level=2)
rule("Commendation", "Each side awards one Commendation after every completed Track, regardless of result. The recipient must have participated and survived. A surviving pilot whose unit was destroyed or who ejected remains eligible. The player chooses the recipient.")
rule("Promotion Priority", "When a named slot is vacant, the eligible unnamed pilot with the most Commendations has priority. Break ties by player choice. If no eligible pilot has a Commendation, choose any eligible pilot. Commendations are permanent and are not spent.")
rule("Retirement", "Between Tracks, a player may permanently retire a pilot. Retirement is irreversible and all development is lost. Retiring a named commander grants a free new named commander with the retired unit slot's purchased skills and Edge 1, but no inherited Commendations, advancement, Edge Abilities, or starting 150 SP.")

chapter_band("4  Battlefield Support")
table(["Campaign Size", "Original BSP", "Modern BSP"], [["Star","12","32"],["Binary","24","64"],["Trinary","36","96"]], [2.4,2.2,2.2])
rule("Workbook Mode", "Use either the Original or Modern BSP schedule for the entire campaign. ComStar receives the selected amount; Clan forces receive 0 BSP unless a Track explicitly grants support.")
rule("Persistence", "Battlefield Support Assets are not persistent campaign units. They do not earn Commendations or advancement, and they do not use campaign repair, salvage, or replacement rules. Purchase them separately for each Track; destruction has no campaign effect unless that Track says otherwise.")
rule("Emplacements", "Emplacements are not general-purchase Assets. Use them only when a Track or Option expressly grants them.")

track_sec = doc.add_section(WD_SECTION_START.CONTINUOUS)
track_sec.top_margin = Inches(.78); track_sec.bottom_margin = Inches(.65); track_sec.left_margin = track_sec.right_margin = Inches(.68)
track_sec.header_distance = Inches(.60); track_sec.footer_distance = Inches(.492); set_columns(track_sec, 2, 360)
chapter_band("5  Track Rules")
rule("Track End", "Unless a Track states otherwise, check every Track End condition during the End Phase. An objective ends a Track only when its rules specifically say so.")
rule("Published Values", "Retain every published objective value that this document does not expressly replace. Ignore all Orders objectives.")
rule("Recovery", "Apply the campaign-wide original-owner recovery rules even when the Draconis Reach version of a Track prohibits salvage.")

track("Assault", [
    ("Conquer", "Completing Conquer does not end the Track. If both sides complete it in the same End Phase, neither side scores it."),
    ("Track End", "During the End Phase, the Track ends if either side has no non-crippled units in play or the turn limit has been reached: Turn 10, or Turn 8 in the shortened format."),
])
track("Breakthrough", [
    ("Objective Sequence", "Completing either preliminary objective does not end the Track."),
    ("Track End", "During the End Phase, the Track ends if one side completes Hold the Field, either side has no units in play, or the turn limit has been reached: Turn 10, or Turn 8 in the shortened format."),
])
track("Meeting Engagement", [
    ("Make Their Acquaintance", "The first side to cripple or destroy one-third of the opposing force's BV scores this objective. If both sides reach the threshold in the same End Phase, neither scores it. Voluntary withdrawals count toward the threshold."),
    ("Force Points Option", "If Force Points are used instead of BV, use a 50% threshold."),
])

track("Flank", [
    ("Defender Emplacements", "The Defender receives two Veteran Medium Emplacements per Campaign Scale. Each emplacement may instead be replaced by up to 20 BSP of emplacement Assets. Place them in the Attacker's half of the battlefield."),
    ("Attacker Entry", "One-third of the Attacker's force enters from its home edge on Turn 1. The remainder enters from a short edge on Turn 2."),
    ("Thresholds", "Crush and Turn the Tide each use a 50% threshold."),
    ("Cut Off Retreat", "Maintain either one non-crippled BattleMech or two Battlefield Support Assets per Campaign Scale near the Defender's home edge for two consecutive End Phases."),
    ("Fall Back", "After Turn the Tide is completed, withdraw at least 50% of the eligible force."),
    ("Ramming Speed", "Add the Ramming Speed objective using its Draconis Reach procedure. It is worth 150 Objective Points."),
    ("Track End", "End when either side has no units in play or at Turn 8 (Turn 6 for the shortened format)."),
])

track("Pushback", [
    ("Vehicle Detachment", "The Attacker receives a separate vehicle-only detachment: 12/24/36 BSP in Original mode or 32/64/96 BSP in Modern mode. This detachment is in addition to normal Battlefield Support."),
    ("Entry", "The primary force enters from its home edge on Turn 1. The vehicle detachment enters from the side edges on Turn 1, halfway toward the Defender's edge."),
    ("Objectives", "Use Push; Crush at one-third; Gutted at 75%; Advance to the Rear instead of Hold Ground; Make Them Hurt at 50%; and the updated Lead objective. Remove Cut Off the Head and ignore Orders."),
    ("Track End", "End when either side has no units in play or at Turn 10 (Turn 8 shortened). After Turn 6, completing Advance to the Rear also ends the Track after the End Phase."),
])
track("Pursuit", [
    ("Roles", "Determine momentum by the campaign roll. The Defender flees and the Attacker pursues; the role assignment is not fixed by faction."),
    ("Entry", "The Defender enters from its home edge on Turn 1. The Attacker enters from the same edge on Turn 2."),
    ("Force Construction", "Use the Tukayyid equal-force construction and campaign recovery rules."),
    ("Fleeing Force", "At least half of the Defender's force must have a maximum movement of 8 or less (16 inches or less in Alpha Strike). Temporary damage does not make a unit qualify."),
    ("Objectives", "Qualifying units count for Escape and Never Here."),
    ("Track End", "End when either side has no non-crippled units in play or at Turn 10."),
])
track("Recon", [
    ("Defender Deployment", "Use the restricted Defender deployment from Draconis Reach."),
    ("Scan", "The Attacker must scan two-thirds of the required targets. Preemptive Strike uses a 25% threshold."),
    ("Observe and Report", "Use Observe and Report in place of Escape. Withdrawal may begin after Turn 4."),
    ("Deny", "Score Deny proportionally. Defender units that withdraw early count as scanned or destroyed, as applicable."),
    ("Track End", "End when either side has no units in play or at Turn 8."),
])
track("Retreat", [
    ("Entry", "The Defender retreats from the Attacker's home edge toward its own. The Attacker enters from that same edge on Turn 2."),
    ("Scan", "Use a two-thirds scan threshold."),
    ("Hammer", "The first side to reach 25% scores Hammer. If both reach it in the same End Phase, neither scores."),
    ("Gauntlet", "Complete Gauntlet when half of the eligible force escapes; crippled units count."),
    ("Track End", "End when either side has no units in play or at Turn 10."),
])
track("Strike", [
    ("Objectives", "Retain the Tukayyid objectives."),
    ("Battlefield", "Use the Draconis Reach building durability and deployment zones. Retain the Tukayyid fortress option."),
    ("Reinforcements", "Replace the secret 1D6+1 arrival with guaranteed Turn 2 flank entry."),
    ("Track End", "End when either side has no units in play or at Turn 10 (Turn 8 shortened)."),
])
track("Supply", [
    ("Track", "Retain the Tukayyid Supply Track. Do not add Objective Raid as a separate Track."),
    ("Warehouses", "Place six buildings and secretly identify two as warehouses. A warehouse must be scanned before loading."),
    ("Components", "Each warehouse holds one component per Campaign Scale, for a total of 2/4/6 components."),
    ("Handling", "Use the Draconis Reach loading, carrying, transfer, and movement-penalty rules for components."),
    ("Objectives", "Retain the Tukayyid casualty and survival objectives."),
    ("Delayed Group", "The delayed group enters on Turn 6."),
    ("Track End", "End when either side has no units in play or at Turn 10."),
    ("Campaign Reward", "Each component carried off the Attacker's home edge awards the Attacker 500 SP. This reward is in addition to the Track's normal Support Allocation."),
])

doc.core_properties.title = "Historical: Battle of Tukayyid - Campaign Rules"
doc.core_properties.subject = "Version 0.1 changes-only campaign playtest"
doc.core_properties.author = "Campaign Development Draft"
OUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(OUT)
print(OUT.resolve())
